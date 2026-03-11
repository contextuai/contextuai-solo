"""
Document Tools for extracting text from PDF and DOCX files.
Provides on-demand document processing for AI chat context.
"""

import logging
from typing import Dict, Any, List
from strands import tool

logger = logging.getLogger(__name__)


class DocumentTools:
    """Tools for extracting text content from PDF and DOCX documents."""

    def __init__(self):
        """Initialize DocumentTools."""
        logger.info("DocumentTools initialized")

    def get_tools(self) -> List:
        """Return list of document tool methods for Strands Agent."""
        return [
            self.extract_pdf_text,
            self.extract_docx_text,
        ]

    @tool
    def extract_pdf_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text content from a PDF file.

        Args:
            file_path: Absolute path to the PDF file

        Returns:
            Dictionary with extracted text, page count, and status
        """
        try:
            import fitz  # PyMuPDF

            logger.info(f"Extracting text from PDF: {file_path}")

            # Open the PDF document
            pdf_doc = fitz.open(file_path)
            page_count = len(pdf_doc)

            # Extract text from all pages
            text_content = []
            for page_num, page in enumerate(pdf_doc, 1):
                page_text = page.get_text()
                if page_text.strip():
                    text_content.append(f"--- Page {page_num} ---\n{page_text}")

            pdf_doc.close()

            full_text = "\n\n".join(text_content)

            logger.info(f"Successfully extracted {len(full_text)} characters from {page_count} pages")

            return {
                "success": True,
                "content": full_text,
                "pages": page_count,
                "characters": len(full_text),
                "file_path": file_path
            }

        except ImportError:
            logger.error("PyMuPDF (fitz) not installed. Install with: pip install PyMuPDF")
            return {
                "success": False,
                "error": "PDF extraction library not available",
                "file_path": file_path
            }
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {e}")
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            }

    @tool
    def extract_docx_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text content from a DOCX file.

        Args:
            file_path: Absolute path to the DOCX file

        Returns:
            Dictionary with extracted text, paragraph count, and status
        """
        try:
            from docx import Document

            logger.info(f"Extracting text from DOCX: {file_path}")

            # Open the DOCX document
            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))

            # Combine paragraph and table text
            full_text = "\n\n".join(paragraphs)
            if table_text:
                full_text += "\n\n--- Tables ---\n" + "\n".join(table_text)

            logger.info(f"Successfully extracted {len(full_text)} characters from {len(paragraphs)} paragraphs")

            return {
                "success": True,
                "content": full_text,
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "characters": len(full_text),
                "file_path": file_path
            }

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return {
                "success": False,
                "error": "DOCX extraction library not available",
                "file_path": file_path
            }
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            return {
                "success": False,
                "error": str(e),
                "file_path": file_path
            }


# Standalone extraction functions for use in ai_chat.py
def extract_pdf_text_sync(file_path: str) -> str:
    """
    Synchronous PDF text extraction for direct use in ai_chat.py.

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text content or error message
    """
    try:
        import fitz

        pdf_doc = fitz.open(file_path)
        text_parts = []

        for page_num, page in enumerate(pdf_doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"[Page {page_num}]\n{page_text}")

        pdf_doc.close()
        return "\n\n".join(text_parts)

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return f"[Error extracting PDF: {e}]"


def extract_docx_text_sync(file_path: str) -> str:
    """
    Synchronous DOCX text extraction for direct use in ai_chat.py.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content or error message
    """
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Include table content
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return "\n\n".join(paragraphs)

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return f"[Error extracting DOCX: {e}]"
